"""
Document processor for Panda.
Processes various document formats and extracts text content.
Optimized for 8GB RAM with async batch processing.
Uses cached tiktoken encoder to reduce memory overhead.
"""

import io
import json
import logging
import os
import asyncio
from typing import List, Optional, Dict, Any, Tuple

import pandas as pd
import numpy as np
from docx import Document as DocxDocument
from PIL import Image
from PyPDF2 import PdfReader

from ..core.model_cache import get_tiktoken_encoder
from ..models.document import Document, DocumentChunk, DocumentMetadata
from ..utils.file_validation import get_safe_filename
from ..utils.memory_monitor import check_memory_available, get_memory_pressure

logger = logging.getLogger(__name__)

# Check OCR availability once at module load time
OCR_AVAILABLE = False
PYTESSERACT = None
try:
    import pytesseract
    PYTESSERACT = pytesseract
    OCR_AVAILABLE = True
    logger.info("OCR (Tesseract) is available")
except ImportError:
    logger.warning("OCR (Tesseract) not available. Image processing disabled.")


class DocumentProcessor:
    """
    Processes various document formats and extracts text content.
    Supports: PDF, DOCX, TXT, CSV, JSON, and images (OCR).
    Optimized for 8GB RAM with batch processing.
    """
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # Use cached tiktoken encoder instead of creating new instance
        self.ocr_available = OCR_AVAILABLE
        self.pytesseract = PYTESSERACT
    
    def _get_tokenizer(self):
        """Get cached tiktoken encoder."""
        return get_tiktoken_encoder("cl100k_base")
    
    async def process_files_batch(
        self,
        files_data: List[Tuple[bytes, str, Optional[str]]]
    ) -> List[Document]:
        """
        Process multiple files in async batches for memory efficiency.
        
        Args:
            files_data: List of (content, filename, file_type) tuples
            
        Returns:
            List of Document objects
        """
        documents = []
        
        # Process files concurrently using asyncio
        tasks = [
            asyncio.create_task(self._process_file_async(content, filename, file_type))
            for content, filename, file_type in files_data
        ]
        
        # Wait for all tasks to complete
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        for result in results:
            if isinstance(result, Exception):
                logger.error(f"Error processing file: {result}")
            elif result is not None:
                documents.append(result)
        
        return documents
    
    async def _process_file_async(
        self,
        file_content: bytes,
        filename: str,
        file_type: Optional[str] = None
    ) -> Document:
        """Async wrapper for process_file."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            lambda: self.process_file(file_content, filename, file_type)
        )
    
    def process_file(
        self, 
        file_content: bytes, 
        filename: str,
        file_type: Optional[str] = None
    ) -> Document:
        """
        Process a file and return a Document object.
        
        Args:
            file_content: Raw bytes of the file
            filename: Original filename (already sanitized)
            file_type: Optional MIME type or extension
            
        Returns:
            Document object with extracted text and chunks
        """
        if not check_memory_available(required_mb=200):
            logger.warning("Low memory detected before processing file")
        
        if file_type is None:
            from ..utils.file_validation import detect_file_type_from_content
            file_type = detect_file_type_from_content(file_content[:8192])
            if file_type is None:
                file_type = "txt"
        
        metadata = DocumentMetadata(
            source_file=get_safe_filename(filename),
            file_type=file_type,
            size_bytes=len(file_content)
        )
        
        document = Document(
            filename=get_safe_filename(filename),
            file_type=file_type,
            size_bytes=len(file_content),
            metadata=metadata,
            processing_status="processing"
        )
        
        try:
            if file_type == "pdf":
                text, meta = self._extract_pdf(file_content)
            elif file_type == "docx":
                text, meta = self._extract_docx(file_content)
            elif file_type == "txt":
                text, meta = self._extract_txt(file_content)
            elif file_type == "csv":
                text, meta = self._extract_csv_streaming(file_content)
            elif file_type == "json":
                text, meta = self._extract_json(file_content)
            elif file_type in ["png", "jpg", "jpeg", "tiff", "bmp"]:
                text, meta = self._extract_image(file_content, filename)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            document.content = text
            document.metadata.total_pages = meta.get("total_pages")
            document.metadata.author = meta.get("author")
            
            document.chunks = self._create_chunks(text, document.document_id, document.metadata)
            document.processing_status = "completed"
            
            logger.info(
                f"Processed {filename}: {len(text)} chars, "
                f"{len(document.chunks)} chunks, {document.get_total_tokens()} tokens"
            )
            
        except Exception as e:
            document.processing_status = "error"
            document.error_message = str(e)
            logger.error(f"Error processing {filename}: {e}")
            raise
        
        return document
    
    def _extract_pdf(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF with memory-efficient batch processing."""
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        meta = {
            "total_pages": len(reader.pages),
            "author": None
        }
        
        if reader.metadata:
            meta["author"] = reader.metadata.get("/Author") or reader.metadata.get("/Creator")
        
        # Process pages in batches for memory efficiency
        batch_size = 3
        total_pages = len(reader.pages)
        
        for batch_start in range(0, total_pages, batch_size):
            batch_end = min(batch_start + batch_size, total_pages)
            
            for i in range(batch_start, batch_end):
                page = reader.pages[i]
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"\n--- Page {i + 1} ---\n{page_text}")
            
            # Check memory pressure between batches
            if batch_end < total_pages and get_memory_pressure() > 80:
                logger.warning(f"High memory pressure during PDF processing, pausing...")
                import time
                time.sleep(0.1)
        
        return "\n".join(text_parts), meta
    
    def _extract_docx(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX."""
        doc_file = io.BytesIO(content)
        doc = DocxDocument(doc_file)
        
        text_parts = []
        meta = {
            "total_pages": None,
            "author": None
        }
        
        if doc.core_properties:
            meta["author"] = doc.core_properties.author
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts), meta
    
    def _extract_txt(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file."""
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                text = content.decode(encoding)
                return text, {"total_pages": None, "author": None}
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
    
    def _extract_csv_streaming(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from CSV file using streaming."""
        csv_file = io.BytesIO(content)
        
        df = None
        for encoding in ["utf-8", "latin-1"]:
            try:
                df = pd.read_csv(csv_file, encoding=encoding, nrows=100)
                break
            except UnicodeDecodeError:
                csv_file.seek(0)
                continue
        
        if df is None:
            raise ValueError("Could not decode CSV file")
        
        csv_file.seek(0)
        total_rows = sum(1 for _ in csv_file) - 1
        
        text_parts = [
            f"CSV Columns: {', '.join(df.columns.tolist())}",
            f"Total Rows: {total_rows}",
            f"Preview (first {len(df)} rows):\n"
        ]
        
        text_parts.append(df.to_string(index=False))
        
        if total_rows > 100:
            text_parts.append(f"\n... ({total_rows - 100} more rows not shown)")
        
        return "\n".join(text_parts), {"total_pages": None, "author": None}
    
    def _extract_json(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from JSON file."""
        for encoding in ["utf-8", "latin-1"]:
            try:
                data = json.loads(content.decode(encoding))
                break
            except UnicodeDecodeError:
                continue
        else:
            raise ValueError("Could not decode JSON file")
        
        text = json.dumps(data, indent=2, ensure_ascii=False)
        return text, {"total_pages": None, "author": None}
    
    def _extract_image(self, content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from image using OCR."""
        if not self.ocr_available:
            raise ValueError("OCR not available. Install pytesseract and Tesseract OCR.")
        
        try:
            image = Image.open(io.BytesIO(content))
            
            if image.mode != "RGB":
                image = image.convert("RGB")
            
            text = self.pytesseract.image_to_string(image)
            
            meta = {
                "total_pages": 1,
                "author": None
            }
            
            return text, meta
            
        except Exception as e:
            raise ValueError(f"OCR failed: {str(e)}")
    
    def _create_chunks(
        self, 
        text: str, 
        document_id: str, 
        metadata: DocumentMetadata
    ) -> List[DocumentChunk]:
        """
        Create overlapping chunks from text using token-based chunking.
        Uses precise token boundaries to prevent token splitting.
        """
        if not text.strip():
            return []
        
        tokenizer = self._get_tokenizer()
        tokens = tokenizer.encode(text)
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(tokens):
            # Get chunk tokens
            end = min(start + self.chunk_size, len(tokens))
            chunk_tokens = tokens[start:end]
            
            # Decode back to text
            chunk_text = tokenizer.decode(chunk_tokens)
            
            # Create chunk
            chunk = DocumentChunk(
                document_id=document_id,
                text=chunk_text,
                token_count=len(chunk_tokens),
                metadata=metadata,
                chunk_index=chunk_index,
                total_chunks=0
            )
            
            chunks.append(chunk)
            chunk_index += 1
            
            # Move start position with overlap
            start += self.chunk_size - self.chunk_overlap
            
            # Avoid infinite loop on small texts
            if start >= end:
                break
        
        # Update total_chunks for each chunk
        for chunk in chunks:
            chunk.total_chunks = len(chunks)
        
        return chunks
    
    def get_token_count(self, text: str) -> int:
        """Get accurate token count using cached encoder."""
        if not text:
            return 0
        tokenizer = self._get_tokenizer()
        return len(tokenizer.encode(text))
    
    def get_document_summary(self, document: Document) -> Dict[str, Any]:
        """Get a summary of the processed document."""
        return {
            "document_id": document.document_id,
            "filename": document.filename,
            "file_type": document.file_type,
            "size_mb": round(document.size_bytes / (1024 * 1024), 2),
            "total_chunks": document.get_chunk_count(),
            "total_tokens": document.get_total_tokens(),
            "processing_status": document.processing_status,
            "error": document.error_message
        }
